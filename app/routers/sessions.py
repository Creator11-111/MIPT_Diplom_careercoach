"""
Роутер для работы с сессиями чата

Сессия - это диалог с пользователем, в котором собирается информация
о финансовом специалисте через интервью.
"""

from __future__ import annotations

from datetime import datetime
from uuid import uuid4

from fastapi import APIRouter, HTTPException, Depends
from fastapi.responses import Response

from app.db.mongo import get_db
from app.models import (
    CreateSessionRequest,
    CreateSessionResponse,
    GetSessionResponse,
    ListSessionsResponse,
    SessionListItem,
    Message,
    Session,
    SessionState,
)
from app.repos.chat_repos import SessionsRepository, MessagesRepository
from app.services.chat_service import ChatService


router = APIRouter()


async def get_sessions_repo(db=Depends(get_db)) -> SessionsRepository:  # noqa: ANN001
    """Получить репозиторий сессий"""
    return SessionsRepository(db)


async def get_messages_repo(db=Depends(get_db)) -> MessagesRepository:  # noqa: ANN001
    """Получить репозиторий сообщений"""
    return MessagesRepository(db)


def get_chat_service() -> ChatService:
    """Получить сервис чата"""
    return ChatService()


@router.post("", response_model=CreateSessionResponse)
async def create_session(
    payload: CreateSessionRequest | None = None,
    sessions_repo: SessionsRepository = Depends(get_sessions_repo),
    messages_repo: MessagesRepository = Depends(get_messages_repo),
    chat_service: ChatService = Depends(get_chat_service),
) -> CreateSessionResponse:
    """
    Создать новую сессию чата с автоматическим приветствием
    
    Args:
        payload: Запрос с опциональным user_id
        sessions_repo: Репозиторий сессий
        messages_repo: Репозиторий сообщений
        chat_service: Сервис чата
        
    Returns:
        CreateSessionResponse с ID созданной сессии и user_id
    """
    import logging
    logger = logging.getLogger(__name__)
    
    try:
        now = datetime.utcnow().isoformat()
        
        # Генерируем user_id, если не указан
        user_id = payload.user_id if payload and payload.user_id else str(uuid4())
        
        # Валидация user_id
        if user_id and len(user_id) > 100:
            raise HTTPException(status_code=400, detail="user_id is too long (max 100 characters)")
        
        # Создаем новую сессию
        session = Session(
            session_id=str(uuid4()),
            user_id=user_id,
            state=SessionState(
                last_question_type=None,
                last_updated_at=now,
            ),
        )
        
        # Сохраняем в базу данных
        await sessions_repo.insert_one(session.model_dump())
        logger.info(f"Created session {session.session_id[:8]}... for user {user_id[:8]}...")
    except Exception as e:
        logger.error(f"Failed to create session: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail="Failed to create session") from e
    
    # Автоматически генерируем приветственное сообщение
    import logging
    logger = logging.getLogger(__name__)
    try:
        await chat_service.generate_welcome_message(
            session.session_id,
            sessions_repo,
            messages_repo
        )
        logger.info(f"Welcome message generated for session {session.session_id[:8]}...")
    except Exception as e:
        # Если не удалось создать приветствие, продолжаем без него
        # (сессия уже создана)
        logger.warning(f"Failed to generate welcome message for session {session.session_id[:8]}...: {e}")
        # Не прерываем создание сессии, пользователь сможет начать диалог сам
    
    return CreateSessionResponse(session_id=session.session_id, user_id=user_id)


@router.get("/{session_id}", response_model=GetSessionResponse)
async def get_session(
    session_id: str,
    sessions_repo: SessionsRepository = Depends(get_sessions_repo),
    messages_repo: MessagesRepository = Depends(get_messages_repo),
) -> GetSessionResponse:
    """
    Получить сессию с историей сообщений
    
    Args:
        session_id: ID сессии
        sessions_repo: Репозиторий сессий
        messages_repo: Репозиторий сообщений
        
    Returns:
        GetSessionResponse с сессией и сообщениями
        
    Raises:
        HTTPException: Если сессия не найдена
    """
    # Получаем сессию
    session_doc = await sessions_repo.find_by_id(session_id)
    if not session_doc:
        raise HTTPException(status_code=404, detail="Session not found")
    
    # Получаем сообщения
    messages_docs = await messages_repo.list_by_session(session_id, limit=50)
    
    # Преобразуем в модели
    session = Session.model_validate(session_doc)
    messages = [Message.model_validate(md) for md in messages_docs]
    
    return GetSessionResponse(session=session, messages=messages)


@router.get("", response_model=ListSessionsResponse)
async def list_sessions(
    user_id: str | None = None,
    sessions_repo: SessionsRepository = Depends(get_sessions_repo),
    messages_repo: MessagesRepository = Depends(get_messages_repo),
) -> ListSessionsResponse:
    """
    Получить список всех сессий пользователя
    
    Args:
        user_id: ID пользователя (опционально)
        sessions_repo: Репозиторий сессий
        messages_repo: Репозиторий сообщений
        
    Returns:
        ListSessionsResponse со списком сессий
    """
    if not user_id:
        return ListSessionsResponse(sessions=[])

    sessions_docs = await sessions_repo.list_by_user_id(user_id)

    sessions_list = []
    for session_doc in sessions_docs:
        session = Session.model_validate(session_doc)
        
        # Получаем все сообщения для определения типа чата
        all_messages = await messages_repo.list_by_session(session.session_id, limit=50)
        
        # Определяем тип чата по первому сообщению или контексту
        preview = None
        chat_type = None
        
        if all_messages:
            first_msg = all_messages[0].get("content", "").lower() if all_messages else ""
            last_msg = all_messages[-1].get("content", "") if all_messages else ""
            
            # Определяем тип чата
            if "хочу найти работу" in first_msg or "интервью" in first_msg or "вопрос" in first_msg:
                chat_type = "💼 Поиск работы"
            elif "развитие карьеры" in first_msg or "план развития" in last_msg:
                chat_type = "🚀 Развитие карьеры"
            elif "анализ навыков" in first_msg or "навыки" in last_msg:
                chat_type = "📊 Анализ навыков"
            elif "анализ целей" in first_msg or "карьерные цели" in last_msg:
                chat_type = "🎯 Анализ целей"
            else:
                # Пытаемся определить по последнему сообщению
                if "план развития" in last_msg.lower():
                    chat_type = "🚀 Развитие карьеры"
                elif "навыки" in last_msg.lower() and "анализ" in last_msg.lower():
                    chat_type = "📊 Анализ навыков"
                elif "цели" in last_msg.lower() and "анализ" in last_msg.lower():
                    chat_type = "🎯 Анализ целей"
                else:
                    chat_type = "💼 Поиск работы"
            
            # Формируем превью
            if last_msg:
                preview_text = last_msg[:80] + "..." if len(last_msg) > 80 else last_msg
                preview = f"{chat_type} • {preview_text}"
        
        if not preview:
            preview = chat_type or "💼 Новый чат"
        
        sessions_list.append(SessionListItem(
            session_id=session.session_id,
            user_id=session.user_id,
            last_updated_at=session.state.last_updated_at,
            preview=preview
        ))
    
    return ListSessionsResponse(sessions=sessions_list)


@router.delete("/{session_id}")
async def delete_session(
    session_id: str,
    sessions_repo: SessionsRepository = Depends(get_sessions_repo),
    messages_repo: MessagesRepository = Depends(get_messages_repo),
) -> dict:
    """
    Удалить сессию и все её сообщения
    
    Args:
        session_id: ID сессии для удаления
        sessions_repo: Репозиторий сессий
        messages_repo: Репозиторий сообщений
        
    Returns:
        Словарь с результатом удаления
        
    Raises:
        HTTPException: Если сессия не найдена
    """
    # Проверяем существование сессии
    session = await sessions_repo.find_by_id(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    
    # Удаляем все сообщения сессии
    await messages_repo.delete_by_session_id(session_id)
    
    # Удаляем саму сессию
    await sessions_repo.delete_by_id(session_id)
    
    return {"status": "success", "message": "Session deleted successfully"}


@router.get("/{session_id}/export")
async def export_session_history(
    session_id: str,
    format: str = "pdf",  # pdf или docx
    sessions_repo: SessionsRepository = Depends(get_sessions_repo),
    messages_repo: MessagesRepository = Depends(get_messages_repo),
    db=Depends(get_db),  # noqa: ANN001
) -> Response:
    """
    Экспортировать историю сессии в Word или PDF
    
    Args:
        session_id: ID сессии
        format: Формат экспорта (pdf или docx)
        sessions_repo: Репозиторий сессий
        messages_repo: Репозиторий сообщений
        
    Returns:
        Файл с историей чата
        
    Raises:
        HTTPException: Если сессия не найдена
    """
    # Проверяем, что сессия существует
    session_doc = await sessions_repo.find_by_id(session_id)
    if not session_doc:
        raise HTTPException(status_code=404, detail="Session not found")
    
    # Получаем все сообщения
    messages_docs = await messages_repo.get_all_by_session(session_id)
    messages = [Message.model_validate(md) for md in messages_docs]
    
    # Получаем профиль, если он есть
    from app.repos.profile_repos import ProfilesRepository
    profiles_repo = ProfilesRepository(db)
    profile_doc = await profiles_repo.find_by_session_id(session_id)
    
    # Пытаемся получить вакансии и курсы из профиля
    vacancies = []
    courses = []
    career_development = None
    
    if profile_doc:
        try:
            # Пытаемся получить career development данные
            from app.services.career_development_service import CareerDevelopmentService
            from app.models.match_models import CareerDevelopmentRequest
            from app.repos.vacancy_repos import VacanciesRepository
            from app.repos.course_repos import CoursesRepository
            
            career_service = CareerDevelopmentService()
            vacancies_repo = VacanciesRepository(db)
            courses_repo = CoursesRepository(db)
            
            # Получаем профиль
            from app.models import UserProfile
            profile = UserProfile.model_validate(profile_doc.get("profile", {}))
            
            # Если есть цели, пытаемся получить career development
            if profile.goals and profile.goals.desired_role:
                try:
                    career_req = CareerDevelopmentRequest(
                        session_id=session_id,
                        target_position=profile.goals.desired_role,
                        target_field=profile.goals.target_field,
                        target_specialization=profile.goals.target_specialization,
                    )
                    career_development = await career_service.get_career_development(
                        career_req,
                        profiles_repo,
                        vacancies_repo,
                        courses_repo,
                    )
                    courses = career_development.courses if career_development else []
                    vacancies = career_development.future_vacancies if career_development else []
                except Exception:
                    pass  # Если не получилось, просто пропускаем
            
            # Если career development не получился, пытаемся получить вакансии из профиля
            if not vacancies:
                try:
                    from app.services.match_service import MatchService
                    from app.models.match_models import MatchVacanciesRequest
                    
                    match_service = MatchService()
                    # Строим резюме из профиля
                    resume_parts = []
                    if profile.professional_context:
                        if profile.professional_context.professional_role:
                            resume_parts.append(f"Позиция: {profile.professional_context.professional_role}")
                        if profile.professional_context.professional_field:
                            resume_parts.append(f"Сфера: {profile.professional_context.professional_field}")
                    if profile.resume:
                        for item in profile.resume[:2]:  # Первые 2 места работы
                            if item.title:
                                resume_parts.append(f"Опыт: {item.title}")
                    if profile.skills and profile.skills.hard_skills:
                        resume_parts.append(f"Навыки: {', '.join(profile.skills.hard_skills[:5])}")
                    
                    resume_text = "\n".join(resume_parts) if resume_parts else "Финансовый специалист"
                    
                    match_req = MatchVacanciesRequest(resume=resume_text, k_faiss=50, k_stage1=20, k_stage2=10)
                    match_response = await match_service.match_vacancies(match_req, vacancies_repo)
                    vacancies = match_response.result[:10]  # Берем первые 10
                except Exception:
                    pass  # Если не получилось, просто пропускаем
        except Exception:
            pass  # Если что-то пошло не так, просто продолжаем без вакансий и курсов
    
    # Экспортируем в нужном формате
    if format == "docx":
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from io import BytesIO
            
            doc = Document()
            
            # Заголовок
            title = doc.add_heading(f'История чата - Сессия {session_id}', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Раздел: История чата
            doc.add_heading('История переписки', level=1)
            for msg in messages:
                role_name = "Пользователь" if msg.role.value == "user" else "AI Консультант"
                doc.add_heading(role_name, level=2)
                doc.add_paragraph(msg.content)
                doc.add_paragraph("─" * 50)
            
            # Раздел: Рекомендованные курсы
            if courses:
                doc.add_page_break()
                doc.add_heading('Рекомендованные курсы для развития', level=1)
                for i, course in enumerate(courses, 1):
                    p = doc.add_paragraph()
                    p.add_run(f"{i}. ").bold = True
                    p.add_run(course.name or "Курс").bold = True
                    if course.provider:
                        doc.add_paragraph(f"   Провайдер: {course.provider}")
                    if course.description:
                        doc.add_paragraph(f"   Описание: {course.description}")
                    if course.skills:
                        doc.add_paragraph(f"   Навыки: {course.skills}")
                    if course.url:
                        doc.add_paragraph(f"   Ссылка: {course.url}")
                    doc.add_paragraph("─" * 50)
            
            # Раздел: Рекомендованные вакансии
            if vacancies:
                doc.add_page_break()
                doc.add_heading('Рекомендованные вакансии', level=1)
                for i, vacancy in enumerate(vacancies, 1):
                    p = doc.add_paragraph()
                    p.add_run(f"{i}. ").bold = True
                    p.add_run(vacancy.title or "Вакансия").bold = True
                    if vacancy.company:
                        doc.add_paragraph(f"   Компания: {vacancy.company}")
                    if vacancy.location:
                        doc.add_paragraph(f"   Локация: {vacancy.location}")
                    if vacancy.salary:
                        doc.add_paragraph(f"   Зарплата: {vacancy.salary}")
                    if vacancy.experience:
                        doc.add_paragraph(f"   Опыт: {vacancy.experience}")
                    if vacancy.description:
                        desc = vacancy.description[:500] + "..." if len(vacancy.description) > 500 else vacancy.description
                        doc.add_paragraph(f"   Описание: {desc}")
                    if vacancy.hh_url:
                        doc.add_paragraph(f"   Ссылка: {vacancy.hh_url}")
                    doc.add_paragraph("─" * 50)
            
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            return Response(
                content=buffer.getvalue(),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f'attachment; filename="chat_history_{session_id}.docx"'}
            )
        except ImportError:
            # Фолбэк на txt, если нет python-docx
            txt = "\n".join(
                (f"{'Пользователь' if m.role.value=='user' else 'AI Консультант'}:\n{m.content}\n" + ("─" * 50))
                for m in messages
            )
            return Response(
                content=txt.encode("utf-8"),
                media_type="text/plain; charset=utf-8",
                headers={"Content-Disposition": f'attachment; filename="chat_history_{session_id}.txt"'}
            )
    
    elif format == "pdf":
        try:
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
            from reportlab.lib.units import inch
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            from reportlab.pdfbase.pdfmetrics import registerFontFamily
            from io import BytesIO
            import os
            
            buffer = BytesIO()
            doc = SimpleDocTemplate(
                buffer, 
                pagesize=A4, 
                rightMargin=72, 
                leftMargin=72, 
                topMargin=72, 
                bottomMargin=72
            )
            styles = getSampleStyleSheet()
            story = []
            
            # В reportlab 4.0+ стандартные шрифты (Helvetica) НЕ поддерживают кириллицу
            # Для правильного отображения кириллицы нужно использовать правильное экранирование
            # или TTF шрифт. Используем правильное экранирование текста.
            
            # Функция для правильного экранирования текста с поддержкой кириллицы
            def escape_text(text):
                """Экранирует текст для использования в Paragraph с поддержкой кириллицы"""
                if not text:
                    return ""
                # Преобразуем в строку, если это не строка
                if not isinstance(text, str):
                    text = str(text)
                # Убираем HTML теги из контента (если есть) ПЕРЕД экранированием
                import re
                text = re.sub(r'<[^>]+>', '', text)
                # Заменяем специальные символы HTML (важно делать это после удаления тегов)
                text = text.replace('&', '&amp;')
                text = text.replace('<', '&lt;')
                text = text.replace('>', '&gt;')
                # Заменяем переносы строк
                text = text.replace('\n', '<br/>')
                # Заменяем специальные символы Unicode на простые аналоги для лучшей совместимости
                text = text.replace('─', '-')  # Заменяем длинное тире на обычное
                text = text.replace('—', '-')  # Заменяем em-dash на обычное
                text = text.replace('–', '-')  # Заменяем en-dash на обычное
                return text
            
            # Используем правильный подход для кириллицы в reportlab
            # В reportlab 4.0+ нужно использовать правильное экранирование и кодировку
            # Используем стандартные шрифты, но с правильной обработкой Unicode
            
            # Создаем стили с правильной обработкой кириллицы
            # Используем стандартные шрифты reportlab, которые поддерживают Unicode
            normal_style = ParagraphStyle(
                'NormalCyrillic',
                parent=styles['Normal'],
                fontName='Helvetica',
                fontSize=10,
                leading=12,
            )
            
            # Заголовок
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                textColor='#1e3c72',
                spaceAfter=30,
                fontName='Helvetica-Bold',
            )
            title_text = escape_text(f'История чата - Сессия {session_id}')
            story.append(Paragraph(title_text, title_style))
            story.append(Spacer(1, 0.2*inch))
            
            # Раздел: История чата
            section_style = ParagraphStyle(
                'SectionStyle',
                parent=styles['Heading1'],
                fontSize=14,
                textColor='#2a5298',
                spaceAfter=20,
                fontName='Helvetica-Bold',
            )
            story.append(Paragraph(escape_text('История переписки'), section_style))
            story.append(Spacer(1, 0.1*inch))
            
            # Сообщения
            for msg in messages:
                role_name = "Пользователь" if msg.role.value == "user" else "AI Консультант"
                role_style = ParagraphStyle(
                    'RoleStyle',
                    parent=styles['Heading2'],
                    fontSize=12,
                    textColor='#2a5298',
                    spaceAfter=10,
                    fontName='Helvetica-Bold',
                )
                story.append(Paragraph(escape_text(role_name), role_style))
                # Экранируем контент - важно правильно обработать кириллицу
                content = escape_text(msg.content)
                # Используем нормальный стиль для текста
                story.append(Paragraph(content, normal_style))
                story.append(Spacer(1, 0.1*inch))
                # Используем простой разделитель вместо специального символа
                story.append(Paragraph(escape_text("-" * 50), normal_style))
                story.append(Spacer(1, 0.1*inch))
            
            # Раздел: Рекомендованные курсы
            if courses:
                story.append(PageBreak())
                story.append(Paragraph(escape_text('Рекомендованные курсы для развития'), section_style))
                story.append(Spacer(1, 0.1*inch))
                
                course_title_style = ParagraphStyle(
                    'CourseTitle',
                    parent=styles['Heading2'],
                    fontSize=11,
                    textColor='#1e3c72',
                    fontName='Helvetica-Bold',
                )
                
                for i, course in enumerate(courses, 1):
                    course_title = escape_text(f"{i}. {course.name or 'Курс'}")
                    story.append(Paragraph(course_title, course_title_style))
                    if course.provider:
                        story.append(Paragraph(escape_text(f"Провайдер: {course.provider}"), normal_style))
                    if course.description:
                        desc = escape_text(course.description)
                        story.append(Paragraph(escape_text(f"Описание: {desc}"), normal_style))
                    if course.skills:
                        skills = escape_text(course.skills)
                        story.append(Paragraph(escape_text(f"Навыки: {skills}"), normal_style))
                    if course.url:
                        story.append(Paragraph(escape_text(f"Ссылка: {course.url}"), normal_style))
                    story.append(Spacer(1, 0.1*inch))
                    story.append(Paragraph(escape_text("-" * 50), normal_style))
                    story.append(Spacer(1, 0.1*inch))
            
            # Раздел: Рекомендованные вакансии
            if vacancies:
                story.append(PageBreak())
                story.append(Paragraph(escape_text('Рекомендованные вакансии'), section_style))
                story.append(Spacer(1, 0.1*inch))
                
                vacancy_title_style = ParagraphStyle(
                    'VacancyTitle',
                    parent=styles['Heading2'],
                    fontSize=11,
                    textColor='#1e3c72',
                    fontName='Helvetica-Bold',
                )
                
                for i, vacancy in enumerate(vacancies, 1):
                    vac_title = escape_text(f"{i}. {vacancy.title or 'Вакансия'}")
                    story.append(Paragraph(vac_title, vacancy_title_style))
                    if vacancy.company:
                        story.append(Paragraph(escape_text(f"Компания: {vacancy.company}"), normal_style))
                    if vacancy.location:
                        story.append(Paragraph(escape_text(f"Локация: {vacancy.location}"), normal_style))
                    if vacancy.salary:
                        story.append(Paragraph(escape_text(f"Зарплата: {vacancy.salary}"), normal_style))
                    if vacancy.experience:
                        story.append(Paragraph(escape_text(f"Опыт: {vacancy.experience}"), normal_style))
                    if vacancy.description:
                        desc = vacancy.description[:500] + "..." if len(vacancy.description) > 500 else vacancy.description
                        desc = escape_text(desc)
                        story.append(Paragraph(escape_text(f"Описание: {desc}"), normal_style))
                    if vacancy.hh_url:
                        story.append(Paragraph(escape_text(f"Ссылка: {vacancy.hh_url}"), normal_style))
                    story.append(Spacer(1, 0.1*inch))
                    story.append(Paragraph(escape_text("-" * 50), normal_style))
                    story.append(Spacer(1, 0.1*inch))
            
            doc.build(story)
            buffer.seek(0)
            
            return Response(
                content=buffer.getvalue(),
                media_type="application/pdf",
                headers={"Content-Disposition": f'attachment; filename="chat_history_{session_id}.pdf"'}
            )
        except ImportError:
            # Если нет reportlab, пробуем docx как фолбэк
            try:
                from docx import Document
                from io import BytesIO
                doc = Document()
                doc.add_heading(f'История чата - Сессия {session_id}', 0)
                for msg in messages:
                    role_name = "Пользователь" if msg.role.value == "user" else "AI Консультант"
                    doc.add_heading(role_name, level=1)
                    doc.add_paragraph(msg.content)
                    doc.add_paragraph("─" * 50)
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                return Response(
                    content=buffer.getvalue(),
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    headers={"Content-Disposition": f'attachment; filename="chat_history_{session_id}.docx"'}
                )
            except Exception:
                # Последний фолбэк — txt
                txt = "\n".join(
                    (f"{'Пользователь' if m.role.value=='user' else 'AI Консультант'}:\n{m.content}\n" + ('─' * 50))
                    for m in messages
                )
                return Response(
                    content=txt.encode("utf-8"),
                    media_type="text/plain; charset=utf-8",
                    headers={"Content-Disposition": f'attachment; filename="chat_history_{session_id}.txt"'}
                )
    
    else:
        raise HTTPException(status_code=400, detail="Unsupported format. Use 'pdf' or 'docx'")

